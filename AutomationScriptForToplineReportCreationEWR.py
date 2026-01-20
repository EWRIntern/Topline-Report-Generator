# -*- coding: utf-8 -*-
"""
Created on Wed Sep 28 07:54:49 2025

@author: Raghavi
"""

# Importing necessary libraries
import streamlit as st
import re
import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL

st.title("Topline Report Generator") # Title page for the UI

# Ask the user for input file which is the crosstab excel file 
uploaded_file = st.file_uploader("Upload a crosstab file", type="xlsx")

# Ask the user for the survey header that needs to be in the topline document
survey_header_name = st.text_input("Enter the survey name exactly as you'd like it to be displayed in the header.")

# Ask the user for the name of the document itself
survey_file_name = st.text_input("Enter the name of your file.")

user_word_by = st.text_input("When you finish creating the crosstabs, indicate the basis used at the end of the sentence. For example, enter banner if it’s by banner, or section 2 if it's by section 2, and so on").strip()

# Check if a file has been uploaded
if uploaded_file and survey_header_name and survey_file_name:
    # Load the workbook
    wb = openpyxl.load_workbook(uploaded_file)
    
    # Load the Tables sheet in which the crosstabs are present
    ws = wb['Tables']
    
    # Get the maximum number of rows and columns in the excel sheet uploaded by the user
    # This is done to determine the sheet dimensions
    m = ws.max_row + 1
    n = ws.max_column + 1
    
    # Get the location of the phrase "Back to TOC in order to identify the start and the end of the questions"
    # This information is stored in toc_locs as a list
    toc_locs = []
    for i in range(1,m):
        if ws[f'A{i}'].value=='Back to TOC':
            toc_locs.append(i)
            
    # Alphabet labels for indexing        
    alphabets = list('abcdefghijklmnopqrstuvwxyz')
    
    # Helper function: replace keywords in column names
    def replace_keywords(column_name):
        replacements = {
            'somewhat more likely': 'SML',
            'somewhat less likely': 'SLL',
            'somewhat': 'Some',
            'favorable': 'Fav',
            'unfavorable': 'Unfav',
            'extremely': 'Ext',
            'concerned': 'Con',
            'convincing': 'Conv',
            'important': 'Imp',
            'unimportant': 'Unimp',
            'satisfied': 'Sat',
            'satisfaction': 'Sat',
            'dissatisfied': 'Dissat',
            'dissatisfacton': 'Dissat',
            'difficult': 'Diff',
            'comfortable': 'Com',
            'uncomfortable': 'Uncom',
            'completely': 'Comp',
            "don't do": 'DNDO',
            "don't know": 'DNK',
            "do not do": 'DNDO',
            "do not know": 'DNK',
            'much more likely': 'MML',
            'much less likely': 'MLL',
            'middle': 'mid'
            }
        
        column_name = column_name.lower()
        for key, value in replacements.items():
            column_name = column_name.replace(key, value)
        column_name = ' '.join(word.capitalize() if "'" not in word else word.capitalize() for word in column_name.split())
        # Capitalize fully specific abbreviations
        for abbr in ['Dndo', 'Dnk', 'Sml', 'Sll', 'Mll', 'Mml']:
            column_name = column_name.replace(abbr, abbr.upper())
            
        return column_name
    
    
    def format_total(x):
        if pd.isna(x):
            return ""
        val = x * 100  # convert to %
        if val == 0:
            return "0"
        elif val < 0.5:
            return "*"
        else:
            return str(int(round(val)))
    
    # Helper function: convert Excel range to DataFrame
    def range_to_df(ws, remove_nan=True, range_address=None):
        # Read the cell values into a list of lists
        data_rows = []
        for row in ws:
            data_cols = []
            for cell in row:
                data_cols.append(cell.value)
            data_rows.append(data_cols)
        df = pd.DataFrame(data_rows[1:])
        df.columns = data_rows[0]
        if remove_nan:
            df.dropna(axis=1, how='all', inplace=True)
    
        #if len(df.columns)==2:
         #   df = df[df.Total>0.0001]
            
        #df['Total'] = df['Total'].map(lambda x: '{:.0f}'.format(x * 100))
           
        df['Total'] = df['Total'].map(format_total)
        
        if len(df.columns)>2 and df.columns[1]!='Total' and df.columns[2]=='Total':
            df.columns = ['Group', 'Statement', 'Total']
            temp = list(df.Statement.unique()[:-1])
            df['Group'] = df['Group'].ffill()
            temp2 = list(df.Group.unique())
            df = df.iloc[:-1,:].pivot(index='Statement', columns='Group', values='Total')
            df = df.loc[temp]
            df = df[temp2]
            
        else:
            df.fillna('', inplace=True)
            temp_cols = []
            for i in df.columns:
                if i=='Total':
                    temp_cols.append(i)
                else:
                    temp_cols.append('')
            df.columns=temp_cols
            if df.iloc[-1, 0] == 'Column Sample Size':
                df = df.iloc[:-1,]
    
        if 'int' in str(df.index.dtype):
            df = df[~df.iloc[:,0].str.contains('NET:')]
            df.Total = df.Total.replace('0%','*%')
            df.columns = ['','']
            if df.shape[0] < 1 or df.shape[1] < 2:
                print(f"⚠️ range_to_df: empty or too small after NET: for {range_address!r}")
                return df
            df.iloc[0, 1] = df.iloc[0, 1] + '%'

            if df.shape[0] > 1: # If dataframe has atleast one row
                df.columns = df.iloc[0, :]
                df = df.iloc[1:, :]
            
        if 'int' not in str(df.index.dtype):
            df = df[~df.index.str.contains('NET:')]
    
        if df.index.name=='Statement':
            df=df.T
            index_temp = []
            for i in range(len(df.index)):
                clean_text = re.sub(r'^[A-Z]\.\s*', '', str(df.index[i]).strip())
                index_temp.append(f'[]{alphabets[i]}. {clean_text}')
    
            df.index = index_temp
            if df.shape[0] < 1 or df.shape[1] < 1:
                print(f"⚠️ range_to_df: empty after NET‑filter/pivot for range {range_address!r}")
                return df
            df.iloc[:, 0] = df.iloc[:, 0] + '%' # First column to percentages
            df.columns = [replace_keywords(col) for col in df.columns]    
        return df
    
    def get_cell_coordinate(row, column):
        column_letter = get_column_letter(column)
        cell_coordinate = f"{column_letter}{row}"
        
        return cell_coordinate
    
    def find_total_col(ws, header_row, n):
        """Return the column index where a cell equals 'Total' on header_row, else None."""
        for col in ws.iter_cols(min_row=header_row, max_row=header_row, max_col=n):
            if any(cell.value == 'Total' for cell in col):
                return next(cell.column for cell in col if cell.value == 'Total')
        return None
    
    def find_table(ws, start_row, max_col, next_toc):
        """
        Detect table end scanning down column max_col starting at start_row.
        Returns (max_row, df) if non-empty df found, else (None, None).
        """
        ct = 0
        max_row = None
        for row in ws.iter_rows(min_col=max_col, max_col=max_col, min_row=start_row, max_row=next_toc):
            cell = row[0]
            if cell.value is None:
                ct += 1
            else:
                ct = 0
            if ct == 2:
                max_row = cell.row - 2
                break
    
        if max_row:
            max_cell = get_cell_coordinate(max_row, max_col)
            rng = f"A{start_row}:{max_cell}"
            df = range_to_df(ws[f'A{start_row}':max_cell], range_address=rng)
            if not df.empty and df.shape[1] > 0:
                return max_row, df
        return None, None

    data=[]
    for i in range(1,m):
        if i in toc_locs:
            by_banner_check = ws[f'A{i+1}'].value
            
            if not isinstance(by_banner_check, str):
                # Skip if the banner cell isn't a string
                continue
            
            pattern = rf"by[\s\.]+.*?{re.escape(user_word_by)}"
            matches = list(re.finditer(pattern, by_banner_check, flags=re.IGNORECASE))
            if not matches:
                continue
            
            last = matches[-1]
            title = by_banner_check[: last.start() ].strip()
            data.append(title)
            
            # --- Locate Total column with fallback (i+3 then i+4) ---
            max_col = None
            for header_row in (i+3, i+4):
                max_col = find_total_col(ws, header_row, n)
                if max_col:
                    header_start_row = header_row  # Remember which header row we locked onto
                    break
            
            if not max_col:
                data.pop()
                print(f"Skipping chunk {i}, no 'Total' column found on rows {i+3} or {i+4}.")
                continue
            
            # --- Determine search boundary for this block ---
            idx = toc_locs.index(i)
            next_toc = toc_locs[idx+1] if idx < len(toc_locs)-1 else m
            
            # --- Find table bounds with fallback start rows (i+3 then i+4) ---
            try:
                max_row, df = find_table(ws, i+3, max_col, next_toc)
            except Exception:
                # fallback
                max_row, df = find_table(ws, i+4, max_col, next_toc)
    
            if df is None:
                data.pop()
                print(f"Skipping chunk {i}, could not locate table bounds starting at rows {i+3} or {i+4}.")
                continue
            
            # --- Keep only non-empty, meaningful tables ---
            if df.empty or df.shape[1] == 0:
                print(f"Skipping chunk starting at row {i+3}, range A{i+3}:{get_cell_coordinate(max_row, max_col)}: no data after NET: filtering")
                print(f"⚠️ Dropping question “{title}” (empty table).")
                data.pop()
                continue
            
            data.append(df)
            
    temp = [isinstance(x, str) and len(x) > 2 and x[0] == 'B' and x[2] == '.' for x in data]
    
    if True in temp:
        data = data[:temp.index(True)]    
        
    def create_element(name):
        return OxmlElement(name)
    
    def create_attribute(element, name, value):
        element.set(qn(name), value)
    
    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
    
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
    
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def add_header_with_page_number(doc):
        section = doc.sections[0]
        header = section.header
    
        paragraph = header.paragraphs[0]
    
        # Add a tab stop to the paragraph
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(doc.sections[0].page_width - doc.sections[0].right_margin, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    
        run = paragraph.add_run()
        run.add_tab()  # Add a tab character
        run.add_tab()
        run.add_text("Page ")
        add_page_number(run)

    def set_table_spacing(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    p_format = paragraph.paragraph_format
                    p_format.space_before = Pt(0)
                    p_format.space_after = Pt(0)
                    
    def set_keep(paragraph, keep_with_next=True, keep_together=True):
        pf = paragraph.paragraph_format
        pf.keep_with_next = keep_with_next
        pf.keep_together = keep_together
        
    def keep_table_together(tbl):
        # Prevent row splits and chain rows so Word moves the whole table if needed
        for r_i, row in enumerate(tbl.rows):
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)
            for p in [p for cell in row.cells for p in cell.paragraphs]:
                # keep each row’s paragraphs together; chain to next row except the last
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = (r_i < len(tbl.rows) - 1)
                
    def set_header_repeat_and_no_row_split(tbl):
        # Mark first row as a repeating header and prevent it from splitting
        hdr = tbl.rows[0]._tr
        trPr = hdr.get_or_add_trPr()
        trPr.append(OxmlElement('w:tblHeader'))
        trPr.append(OxmlElement('w:cantSplit'))
    
        # Prevent splitting for all other rows, but allow page breaks BETWEEN rows
        for r in tbl.rows[1:]:
            tr = r._tr
            trPr = tr.get_or_add_trPr()
            trPr.append(OxmlElement('w:cantSplit'))
    
        # Ensure paragraphs don't chain to the next row (so Word may break BETWEEN rows)
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = False   # allow break after the row
                    p.paragraph_format.keep_together = True     # keep lines within the cell together

    # Create a new Word document
    doc = Document()
    
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    section.different_first_page_header_footer = True
     
    # Selecting the header
    header = section.first_page_header
     
    # Selecting the paragraph already present in the header section
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
     
    # Adding the centred zoned header
    logo_run = header_para.add_run()
    logo_run.add_picture("EWR_Logo.png", width=Inches(2))
    
    header.add_paragraph().add_run()#.add_break()
    
    # Create a header object
    header2 = section.header
    
    # Add text to the header
    header_text = header2.paragraphs[0]
    header_text.text = survey_header_name
    header_text.style.font.name = 'Chaparral Semibold'  # Font name
    header_text.style.font.size = Pt(12) 
    
    add_header_with_page_number(doc)
    
    # Add content to the first page
    for i in data:
        if isinstance(i, str):
            # If the element is a string, write it to the document
            modified_text = i #.replace('.', ' ')
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(modified_text)
            font = run.font
            font.name = 'Acumin Pro'  # Change font style as needed
            font.size = Pt(12) 
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            set_keep(paragraph, keep_with_next=True, keep_together=True)
            blank = doc.add_paragraph()
            blank.paragraph_format.space_before = Pt(0)
            blank.paragraph_format.space_after = Pt(0)
            set_keep(blank, keep_with_next=True, keep_together=True)
            rblank = blank.add_run(" ")      # ensure the paragraph has a run so size applies
            rblank.font.name = 'Acumin Pro'
            rblank.font.size = Pt(12)
            
        # LEFT HERE LAST TIME    
        elif isinstance(i, pd.DataFrame):
            # If the element is a DataFrame, assume it's a table and add it to the document
            table = doc.add_table(rows=len(i) + 1, cols=len(i.columns) + 1)  # Add one extra column for the index
            table.autofit = True  # Disable autofitting to prevent loss of margins
            table.alignment = 1  # Set alignment to center
            is_matrix = len(i.columns) > 2
            if is_matrix:
                set_header_repeat_and_no_row_split(table)
            if not is_matrix:
                pass
    
            # Add index values
            if i.index.dtype.kind != 'i':
                for row_idx, idx_value in enumerate(i.index, start=1):
                    cell = table.cell(row_idx, 0)
                    cell.text = str(idx_value)
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(12)  # Set font size to 12 points
                    cell_font.name = 'Acumin Pro'
                    cell.width = Inches(10.0)
            
            # Add data column headers
            for col_idx, col_name in enumerate(i.columns):
                cell2 = table.cell(0, col_idx + 1)
                cell2.text = col_name
                # cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                cell_font = cell2.paragraphs[0].runs[0].font
                cell_font.size = Pt(10) if is_matrix else Pt(12)  # Set font size to 12 points
                cell_font.name = 'Acumin Pro'
                cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if is_matrix else WD_ALIGN_PARAGRAPH.LEFT
    
            # Add data rows
            for row_idx, (_, row) in enumerate(i.iterrows(), start=1):
                for col_idx, cell_value in enumerate(row, start=1):
                    cell3 = table.cell(row_idx, col_idx)
                    cell3.text = str(cell_value)
                    # cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                    cell_font = cell3.paragraphs[0].runs[0].font
                    cell_font.size = Pt(12)  # Set font size to 12 points
                    cell_font.name = 'Acumin Pro'
                    if is_matrix:
                        cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            if not is_matrix:
                keep_table_together(table)
                    
            # doc.add_paragraph() # Empty space after every question
            spacer_paragraph = doc.add_paragraph()
            spacer_paragraph_format = spacer_paragraph.paragraph_format
            spacer_paragraph_format.space_after = Pt(0)
            spacer_paragraph_format.space_before = Pt(0)  # Minimum possible spacing
            
            run = spacer_paragraph.add_run(" ")            # add a space so the run exists
            run.font.name = 'Acumin Pro'
            run.font.size = Pt(12)
            
            set_table_spacing(table)

    
    # Save the document
    file_name = f"{survey_file_name}.docx"
    doc.save(file_name)

    with open(file_name, "rb") as file:
        st.download_button(
            label="Download Topline Report",
            data=file,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.error("Please upload an Excel file.")

