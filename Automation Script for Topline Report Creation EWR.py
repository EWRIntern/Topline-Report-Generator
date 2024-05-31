# -*- coding: utf-8 -*-
"""
Created on Thu May 30 09:03:45 2024

@author: raghavi
"""

import streamlit as st
import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

st.title("Topline Report Generator")

uploaded_file = st.file_uploader("Upload a crosstab file", type="xlsx")
survey_name = st.text_input("Enter the survey name exactly as you'd like it to be displayed in the header.")

# Check if a file has been uploaded
if uploaded_file and survey_name:
    # Load the workbook
    wb = openpyxl.load_workbook(uploaded_file)
    
    # Option to select sheet
    #sheet_names = wb.sheetnames
    #selected_sheet = st.selectbox("Select a sheet", sheet_names)
    
    # Load the selected sheet
    #ws = wb[selected_sheet]
    ws = wb['Tables']
    
    m = ws.max_row + 1
    n = ws.max_column + 1
    
    toc_locs = []
    for i in range(1,m):
        if ws[f'A{i}'].value=='Back to TOC':
            toc_locs.append(i)
            
    def replace_keywords(column_name):
        replacements = {
            'somewhat': 'Some',
            'favorable': 'Fav',
            'unfavorable': 'Unfav',
            'extremely': 'Ext',
            'concerned': 'Con',
            'convincing': 'Conv',
            'important': 'Imp',
            'unimportant': 'Unimp',
            'satisfied': 'Sat',
            'satisfaction': 'Sat',
            'dissatisfied': 'Dissat',
            'dissatisfacton': 'Dissat',
            'difficult': 'Diff',
            'comfortable': 'Com',
            'uncomfortable': 'Uncom',
            'completely': 'Comp'
            }
        for key, value in replacements.items():
            column_name = column_name.lower().replace(key, value)
            
        return column_name

    def range_to_df(ws, remove_nan=True):
        # Read the cell values into a list of lists
        data_rows = []
        for row in ws:
            data_cols = []
            for cell in row:
                data_cols.append(cell.value)
            data_rows.append(data_cols)
        df = pd.DataFrame(data_rows[1:])
        df.columns = data_rows[0]
        if remove_nan:
            df.dropna(axis=1, how='all', inplace=True)
        
        df['Total'] = df['Total'].map(lambda x: '{:.0%}'.format(x))
    
        if len(df.columns)>2 and df.columns[1]!='Total'and df.columns[2]=='Total':
            df.columns = ['Group', 'Statement', 'Total']
            temp = list(df.Statement.unique()[:-1])
            df['Group'] = df['Group'].ffill()
            df = df.iloc[:-1,:].pivot(index='Statement', columns='Group', values='Total')
            df = df.loc[temp]
            
        else:
            df.fillna('', inplace=True)
            temp_cols = []
            for i in df.columns:
                if i=='Total':
                    temp_cols.append(i)
                else:
                    temp_cols.append('')
            df.columns=temp_cols
            if df.iloc[-1, 0] == 'Column Sample Size':
                df = df.iloc[:-1,]
    
        if 'int' in str(df.index.dtype):
            df = df[~df.iloc[:,0].str.contains('NET:')]
            
        if 'int' not in str(df.index.dtype):
            df = df[~df.index.str.contains('NET:')]
    
        if df.index.name=='Statement':
            df=df.T
            df.columns = [replace_keywords(col) for col in df.columns]    
            
        return df
    
    def get_cell_coordinate(row, column):
        column_letter = get_column_letter(column)
        cell_coordinate = f"{column_letter}{row}"
        
        return cell_coordinate

    data=[]
    for i in range(1,m):
        if i in toc_locs:
            by_banner_check = ws[f'A{i+1}'].value
            
            if 'by BANNER' in by_banner_check:
                by_ban_ind = by_banner_check.find('by BANNER')
                data.append(by_banner_check[:by_ban_ind])
                
            elif 'by.BANNER' in by_banner_check:
                by_ban_ind = by_banner_check.find('by.BANNER')
                data.append(by_banner_check[:by_ban_ind])
                
            if 'by BANNER' in by_banner_check or 'by.BANNER' in by_banner_check:
                max_col = None
                max_row = None
                for col in ws.iter_cols(min_row=i+3, max_row = i+3, max_col = n):
                    for cell in col:
                        if cell.value=='Total':
                            max_col = cell.column
                            break
                    if cell.value=='Total':
                        break
                if toc_locs.index(i)<len(toc_locs)-1:
                    next_toc_loc = toc_locs[toc_locs.index(i)+1]
                else:
                    next_toc_loc = m
    
                ct=0
                
                for row in ws.iter_rows(min_col=max_col, max_col=max_col, min_row = i+3, max_row = next_toc_loc):
                    for cell in row:
                        if cell.value==None:
                            ct+=1
                            break
                        else:
                            ct=0
                    if ct==2:
                        max_row = cell.row-2
                        break
    
                if (not max_col) or (not max_row):
                    data.pop()
                    print(f'Table format not found for chunk {i}. Skipping....')
                else:
                    max_cell = get_cell_coordinate(max_row, max_col)
                    dataframe = range_to_df(ws[f'A{i+3}':max_cell])
                    data.append(dataframe)
            else:
                continue    
        
    def create_element(name):
        return OxmlElement(name)
    
    def create_attribute(element, name, value):
        element.set(qn(name), value)
    
    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
    
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
    
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def add_header_with_page_number(doc):
        section = doc.sections[0]
        header = section.header
    
        paragraph = header.paragraphs[0]
    
        # Add a tab stop to the paragraph
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(doc.sections[0].page_width - doc.sections[0].right_margin, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    
        run = paragraph.add_run()
        run.add_tab()  # Add a tab character
        run.add_tab()
        run.add_text("Page ")
        add_page_number(run)

    def set_table_spacing(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    p_format = paragraph.paragraph_format
                    p_format.space_before = Pt(0)
                    p_format.space_after = Pt(0)
                    
    # Create a new Word document
    doc = Document()
    
    section = doc.sections[0]
    
    section.different_first_page_header_footer = True
     
    # Selecting the header
    header = section.first_page_header
     
    # Selecting the paragraph already present in the header section
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
     
    # Adding the centred zoned header
    logo_run = header_para.add_run()
    logo_run.add_picture("EWR_Logo.png", width=Inches(2))
    
    header.add_paragraph().add_run()#.add_break()
    
    # Create a header object
    header2 = section.header
    
    # Add text to the header
    header_text = header2.paragraphs[0]
    header_text.text = survey_name
    header_text.style.font.name = 'Chaparral Semibold'  # Font name
    header_text.style.font.size = Pt(12) 
    
    add_header_with_page_number(doc)
    
    # Add content to the first page
    for i in data:
        if isinstance(i, str):
            # If the element is a string, write it to the document
            modified_text = i.replace('.', ' ')
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(modified_text)
            font = run.font
            font.name = 'Acumin Pro'  # Change font style as needed
            font.size = Pt(12) 
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        elif isinstance(i, pd.DataFrame):
            # If the element is a DataFrame, assume it's a table and add it to the document
            table = doc.add_table(rows=len(i) + 1, cols=len(i.columns) + 1)  # Add one extra column for the index
            table.autofit = True  # Disable autofitting to prevent loss of margins
            table.alignment = 1  # Set alignment to center
    
            # Add index values
            if i.index.dtype.kind != 'i':
                for row_idx, idx_value in enumerate(i.index, start=1):
                    cell = table.cell(row_idx, 0)
                    cell.text = str(idx_value)
                    cell_font = cell.paragraphs[0].runs[0].font
                    cell_font.size = Pt(12)  # Set font size to 12 points
                    cell_font.name = 'Acumin Pro'
    
            # Add data column headers
            for col_idx, col_name in enumerate(i.columns):
                cell2 = table.cell(0, col_idx + 1)
                cell2.text = col_name
                cell_font = cell2.paragraphs[0].runs[0].font
                cell_font.size = Pt(12)  # Set font size to 12 points
                cell_font.name = 'Acumin Pro'
    
            # Add data rows
            for row_idx, (_, row) in enumerate(i.iterrows(), start=1):
                for col_idx, cell_value in enumerate(row, start=1):
                    cell3 = table.cell(row_idx, col_idx)
                    cell3.text = str(cell_value)
                    cell_font = cell3.paragraphs[0].runs[0].font
                    cell_font.size = Pt(12)  # Set font size to 12 points
                    cell_font.name = 'Acumin Pro'
                    
            set_table_spacing(table)
    
    # Save the document
    file_name = f"Topline Report - {survey_name}.docx"
    doc.save(file_name)

    with open(file_name, "rb") as file:
        st.download_button(
            label="Download Topline Report",
            data=file,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.error("Please upload an Excel file.")